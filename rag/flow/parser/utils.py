#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

"""解析器工具函数模块。

提供文档解析过程中的通用工具函数，包括：
- 目录（TOC）移除：支持纯文本、PDF、Word 三种场景
- 页眉/页脚移除：支持 DOCX 和 HTML 格式
- Word 文档大纲提取
- 图片/表格区域的视觉模型增强描述
"""

import re
from io import BytesIO

from bs4 import BeautifulSoup
from docx import Document
from api.db.services.llm_service import LLMBundle
from api.db.joint_services.tenant_model_service import (
    get_tenant_default_model_by_type, get_model_config_from_provider_instance,
)
from common.constants import LLMType
from deepdoc.parser.figure_parser import VisionFigureParser
from rag.nlp import is_english, random_choices, remove_contents_table


def remove_toc(items):
    """移除文本段落列表中的目录内容。

    使用 NLP 模块的 remove_contents_table 函数识别并删除目录项。
    适用于纯文本、HTML 等没有结构化大纲信息的场景。

    Args:
        items: 文本段落列表，每项可以是 str、dict（含 "text" 键）或 list

    Returns:
        tuple: (过滤后的段落列表, 保留项的原始索引列表)
    """
    # 为每个段落建立 (文本内容, 原始索引) 的索引列表
    indexed = [(_item_text(item), i) for i, item in enumerate(items)]
    # 调用 NLP 模块去除目录内容（原地修改 indexed）
    remove_contents_table(indexed, eng=_is_english(indexed))
    # 提取保留项的索引
    kept_indices = [i for _, i in indexed]
    return [items[i] for i in kept_indices], kept_indices


def extract_docx_header_footer_texts(filename=None, binary=None):
    """从 DOCX 文件中提取所有页眉和页脚文本。

    遍历文档每个节的页眉和页脚，收集其中的段落文本和表格文本。
    用于后续的页眉/页脚过滤步骤。

    Args:
        filename: DOCX 文件路径（与 binary 二选一）
        binary: DOCX 文件的二进制内容

    Returns:
        set: 页眉/页脚中出现的唯一文本集合（已做空白规范化处理）
    """
    # 根据输入类型加载 DOCX 文档对象
    doc = Document(filename) if binary is None else Document(BytesIO(binary))
    texts = set()
    for section in doc.sections:
        # 遍历当前节的页眉和页脚
        for container in (section.header, section.footer):
            # 提取段落文本
            for paragraph in container.paragraphs:
                normalized = re.sub(r"\s+", " ", paragraph.text).strip()
                if normalized:
                    texts.add(normalized)
            # 提取表格文本
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        normalized = re.sub(r"\s+", " ", cell.text).strip()
                        if normalized:
                            texts.add(normalized)
    return texts


def remove_header_footer_docx_sections(items, header_footer_texts):
    """从 DOCX 解析结果中移除页眉/页脚对应的段落。

    将每个段落的文本与从页眉/页脚中提取的文本集合进行匹配，
    匹配成功的段落被视为重复的页眉/页脚内容并予以移除。

    Args:
        items: DOCX 解析后的段落列表
        header_footer_texts: extract_docx_header_footer_texts 返回的页眉/页脚文本集合

    Returns:
        list: 移除页眉/页脚段落后的结果列表
    """
    if not header_footer_texts:
        return items

    filtered = []
    for item in items:
        text = _item_text(item)
        # 规范化段落文本并与页眉/页脚文本集合比对
        normalized = re.sub(r"\s+", " ", text).strip() if isinstance(text, str) else ""
        if normalized and normalized in header_footer_texts:
            continue  # 匹配到页眉/页脚文本，跳过
        filtered.append(item)
    return filtered


def remove_header_footer_html_blob(blob):
    """移除 HTML 内容中的页眉/页脚元素。

    通过 BeautifulSoup 解析 HTML，查找并移除以下元素：
    - <header> / <footer> 标签
    - role 属性为 "banner" 或 "contentinfo" 的元素

    Args:
        blob: HTML 内容的字节串

    Returns:
        bytes: 移除页眉/页脚后的 HTML 字节串
    """
    soup = BeautifulSoup(blob, "html.parser")
    # 查找并移除所有页眉/页脚相关元素
    for element in soup.find_all(
        lambda tag: tag.name in {"header", "footer"}
        or tag.get("role") in {"banner", "contentinfo"}
    ):
        element.decompose()  # 从 DOM 树中彻底移除元素
    return str(soup).encode("utf-8")


def extract_word_outlines(filename, binary=None):
    """从 Word 文档中提取标题大纲（目录结构）。

    解析文档中所有应用了 "Heading N" 样式的段落，
    提取标题文本和层级，用于生成文档大纲和目录移除。

    Args:
        filename: DOCX 文件路径（与 binary 二选一）
        binary: DOCX 文件的二进制内容

    Returns:
        list[tuple]: 大纲列表，每项为 (标题文本, 层级, None)
                     层级从 0 开始（Heading 1 → 0, Heading 2 → 1, ...）
    """
    doc = Document(filename) if binary is None else Document(BytesIO(binary))
    outlines = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        # 匹配 "Heading 1", "Heading 2" 等样式名
        match = re.search(r"Heading\s*(\d+)", style_name, re.I)
        if not match:
            continue
        # 层级减 1 使其从 0 开始计数
        outlines.append((text, int(match.group(1)) - 1, None))
    return outlines


def remove_toc_pdf(items, outlines):
    """基于 PDF 大纲信息移除目录页。

    在大纲中查找标记为"目录"的条目，确定目录的起始页和结束页，
    然后删除这些页面范围内的所有内容。

    支持识别的目录标题关键词：
    - 中文：目录、目次、致谢
    - 英文：contents、table of contents、acknowledge

    Args:
        items: PDF 解析后的段落列表（每项需含 page_number 字段）
        outlines: PDF 大纲列表，每项为 (标题, 层级, 页码)

    Returns:
        list: 移除目录页后的段落列表
    """
    if not outlines:
        return items

    toc_start_page = None
    content_start_page = None
    # 在大纲中查找目录起始页和正文起始页
    for i, (title, level, page_no) in enumerate(outlines):
        # 匹配目录/致谢标题
        if re.match(r"(contents|目录|目次|table of contents|致谢|acknowledge)$", title.split("@@")[0].strip().lower()):
            toc_start_page = page_no
            # 查找同级的下一个非目录标题作为正文开始页
            for next_title, next_level, next_page_no in outlines[i + 1:]:
                if next_level != level:
                    continue
                if re.match(r"(contents|目录|目次|table of contents|致谢|acknowledge)$", next_title.split("@@")[0].strip().lower()):
                    continue
                content_start_page = next_page_no
                break
            break

    if content_start_page:
        # 保留正文起始页之后的内容
        return [item for item in items if not (toc_start_page <= item["page_number"] < content_start_page)]
    return items


def remove_toc_word(items, outlines):
    """从 Word 文档解析结果中移除目录内容。

    分两步处理：
    1. 如果有大纲信息：利用大纲标题进行更精确的目录识别，删除目录页、点线引导符行等
    2. 最后调用通用的 remove_toc 进行兜底清理

    Args:
        items: Word 文档解析后的段落列表
        outlines: 文档大纲列表，每项为 (标题, 层级, 页码)

    Returns:
        list: 移除目录后的段落列表
    """
    if not outlines:
        # 无大纲时直接使用通用目录移除
        filtered_items, _ = remove_toc(items)
        return filtered_items

    # 构建大纲标题集合，用于精确匹配
    outline_titles = [title.split("@@")[0].strip().lower() for title, _, _ in outlines if title]
    if outline_titles:
        indexed = [(_item_text(item), i) for i, item in enumerate(items)]
        i = 0
        while i < len(indexed):
            # 找到目录起始位置（匹配目录/致谢标题）
            if not re.match(r"(contents|目录|目次|table of contents|致谢|acknowledge)$", indexed[i][0].split("@@")[0].strip().lower()):
                i += 1
                continue
            indexed.pop(i)
            # 删除目录中的所有条目，直到遇到正文标题
            while i < len(indexed):
                text = indexed[i][0]
                normalized = text.split("@@")[0].strip().lower()
                if not normalized:
                    indexed.pop(i)
                    continue
                # 如果匹配到正文大纲标题，停止删除
                if any(normalized.startswith(title) or title.startswith(normalized) for title in outline_titles):
                    indexed.pop(i)
                    continue
                # 如果包含点线引导符（如 "1.1 简介 ...... 5"），视为目录条目
                if re.search(r"(\.{2,}|…{2,}|·{2,}|[ ]{2,})\s*\d+\s*$", text):
                    indexed.pop(i)
                    continue
                break
            break
        items = [items[i] for _, i in indexed]

    # 最后用通用方法做兜底清理
    filtered_items, _ = remove_toc(items)
    return filtered_items


def _item_text(item):
    """从不同类型的段落项中提取文本内容。

    支持三种段落格式：
    - str：直接返回
    - dict：返回 "text" 字段的值
    - list/tuple：返回第一个元素

    Args:
        item: 段落项，可以是 str、dict 或 list

    Returns:
        str: 提取的文本内容
    """
    if isinstance(item, str):
        return item
    if isinstance(item, dict):
        return item["text"]
    return item[0]


def _is_english(indexed):
    """判断索引化的文本列表是否为英文内容。

    从文本列表中随机采样最多 200 条，使用 NLP 语言检测判断。

    Args:
        indexed: [(文本, 索引), ...] 格式的列表

    Returns:
        bool: True 表示英文内容，False 表示非英文
    """
    texts = [text for text, _ in indexed if text]
    if not texts:
        return False
    return is_english(random_choices(texts, k=200))


def enhance_media_sections_with_vision(
    sections,
    tenant_id,
    vlm_conf=None,
    callback=None,
):
    """使用视觉语言模型（VLM）增强图片和表格区域的文本描述。

    对解析结果中标记为 "image" 或 "table" 类型且包含图片数据的段落，
    调用 VisionFigureParser 生成图片/表格的文字描述，并将描述追加到
    段落的 text 字段中。

    这个步骤可以显著提升后续检索阶段对图片/表格内容的召回率。

    Args:
        sections: 解析后的段落列表，每项为包含 image、doc_type_kwd 等字段的 dict
        tenant_id: 租户 ID，用于获取模型配置
        vlm_conf: VLM 模型配置，需包含 llm_id 字段
        callback: 进度回调函数

    Returns:
        list: 增强后的段落列表（原地修改）
    """
    if not sections or not tenant_id:
        return sections

    # 获取视觉语言模型实例
    try:
        try:
            # 优先使用用户指定的 VLM 模型
            vision_model_config = get_model_config_from_provider_instance(
                tenant_id, LLMType.IMAGE2TEXT, vlm_conf["llm_id"]
            )
        except Exception:
            # 降级使用租户默认的 IMAGE2TEXT 模型
            vision_model_config = get_tenant_default_model_by_type(
                tenant_id, LLMType.IMAGE2TEXT
            )
        vision_model = LLMBundle(tenant_id, vision_model_config)
    except Exception:
        # 模型不可用时静默返回，不阻塞解析流程
        return sections

    for item in sections:
        # 仅处理图片和表格类型的段落
        if item.get("doc_type_kwd") not in {"image", "table"}:
            continue
        if item.get("image") is None:
            continue

        text = item.get("text") or ""
        try:
            # 调用 VisionFigureParser 生成图片/表格描述
            parsed = VisionFigureParser(
                vision_model=vision_model,
                figures_data=[((item["image"], [""]), [(0, 0, 0, 0, 0)])],
                context_size=0,
            )(callback=callback)
        except Exception:
            continue

        if not parsed:
            continue

        # VisionFigureParser 返回格式：[((image, text_or_text_list), positions), ...]
        # first_result[0] 是 (image, parsed_text) 元组
        # first_result[0][1] 是解析出的文本描述
        first_result = parsed[0]
        image_and_text = first_result[0]
        parsed_text = str(image_and_text[1] or "").strip()

        if parsed_text:
            # 将视觉模型生成的描述追加到原文本后面
            item["text"] = f"{text}\n{parsed_text}" if text else parsed_text

    return sections
