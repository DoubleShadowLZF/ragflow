#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

"""
Word 文档解析器

基于 python-docx 库解析 .docx 文件，提取段落文本和表格数据。
支持以下功能：
1. 文本提取：按页面范围过滤，处理分页符（lastRenderedPageBreak）
2. 图片提取：从段落中提取嵌入图片，支持损坏图片的容错降级读取
3. 表格提取：智能识别表格结构（数字型/文本型），自动推断表头行，
   并以 "表头: 值" 的格式输出结构化数据
4. 表格内容类型识别：通过正则 + 分词判断单元格的数据类型
   （日期 Dt、数字 Nu、货币 Ca、英文 En、人名 Nr、文本 Tx 等）
"""

from docx import Document
import re
import pandas as pd
from collections import Counter
from rag.nlp import rag_tokenizer
from io import BytesIO
import logging
from common.constants import MAXIMUM_PAGE_NUMBER
from docx.image.exceptions import (
    InvalidImageStreamError,
    UnexpectedEndOfFileError,
    UnrecognizedImageError,
)
from rag.utils.lazy_image import LazyImage

class RAGFlowDocxParser:
    """Word 文档解析器（.docx 格式）

    解析 Word 文档的段落文本、嵌入图片和表格数据。
    支持按页面范围分段提取，表格自动识别表头并结构化输出。
    """

    def get_picture(self, document, paragraph):
        """从段落中提取嵌入的图片

        解析段落的 XML 结构查找 pic:pic 元素，通过 r:embed 引用
        找到对应的图片部件并提取其二进制数据。

        容错策略：
        - 如果图片对象损坏（UnrecognizedImageError 等），降级为 blob 直读
        - 如果 image.blob 不可用，尝试 related_part.blob 回退

        Args:
            document: python-docx Document 对象
            paragraph: 目标段落对象

        Returns:
            LazyImage 对象（包含图片二进制数据的懒加载封装），无图片时返回 None
        """
        imgs = paragraph._element.xpath(".//pic:pic")
        if not imgs:
            return None
        image_blobs = []
        for img in imgs:
            embed = img.xpath(".//a:blip/@r:embed")
            if not embed:
                continue
            embed = embed[0]
            image_blob = None
            try:
                related_part = document.part.related_parts[embed]
            except Exception as e:
                logging.warning(f"Skipping image due to unexpected error getting related_part: {e}")
                continue

            # 尝试获取 image.blob，如果图片损坏则降级
            try:
                image = related_part.image
                if image is not None:
                    image_blob = image.blob
            except (
                UnrecognizedImageError,
                UnexpectedEndOfFileError,
                InvalidImageStreamError,
                UnicodeDecodeError,
            ) as e:
                logging.info(f"Damaged image encountered, attempting blob fallback: {e}")
            except Exception as e:
                logging.warning(f"Unexpected error getting image, attempting blob fallback: {e}")

            if image_blob is None:
                image_blob = getattr(related_part, "blob", None)
            if image_blob:
                image_blobs.append(image_blob)
        if not image_blobs:
            return None
        return LazyImage(image_blobs)


    def __extract_table_content(self, tb):
        """提取表格内容并转为结构化文本

        处理流程：
        1. 将表格转为 pandas DataFrame
        2. 调用 __compose_table_content 进行表头推断和格式化

        Args:
            tb: python-docx Table 对象

        Returns:
            结构化表格文本列表
        """
        df = []
        for row in tb.rows:
            df.append([c.text for c in row.cells])
        return self.__compose_table_content(pd.DataFrame(df))

    def __compose_table_content(self, df):
        """将 DataFrame 格式的表格转为带表头的结构化文本行

        核心逻辑：
        1. 识别每列单元格的数据类型（blockType 函数）
        2. 统计表格中占比最大的数据类型（max_type）
        3. 如果 max_type 是数字型("Nu")，重新识别表头行（数字行之前的行为表头）
        4. 对每行数据，以 "表头1,表头2: 单元格值" 格式输出
        5. 窄表（≤3列）用分号连接，宽表每行独立

        blockType 识别的数据类型：
        - Dt: 日期（如 2023年1月1日、Q1 等）
        - Nu: 纯数字/百分比
        - Ca: 编号/代码（如 ABC123）
        - En: 英文单词
        - Nr: 人名（分词后识别为 nr 标签的单字）
        - Tx: 短文本（3-12 个 token）
        - Lx: 长文本（>12 个 token）
        - Sg: 单字符
        - Ot: 其他
        """

        def blockType(b):
            """判断单元格的数据类型"""
            pattern = [
                ("^(20|19)[0-9]{2}[年/-][0-9]{1,2}[月/-][0-9]{1,2}日*$", "Dt"),
                (r"^(20|19)[0-9]{2}年$", "Dt"),
                (r"^(20|19)[0-9]{2}[年/-][0-9]{1,2}月*$", "Dt"),
                ("^[0-9]{1,2}[月/-][0-9]{1,2}日*$", "Dt"),
                (r"^第*[一二三四1-4]季度$", "Dt"),
                (r"^(20|19)[0-9]{2}年*[一二三四1-4]季度$", "Dt"),
                (r"^(20|19)[0-9]{2}[ABCDE]$", "DT"),
                ("^[0-9.,+%/ -]+$", "Nu"),
                (r"^[0-9A-Z/\._~-]+$", "Ca"),
                (r"^[A-Z]*[a-z' -]+$", "En"),
                (r"^[0-9.,+-]+[0-9A-Za-z/$￥%<>（）()' -]+$", "NE"),
                (r"^.{1}$", "Sg")
            ]
            for p, n in pattern:
                if re.search(p, b):
                    return n
            tks = [t for t in rag_tokenizer.tokenize(b).split() if len(t) > 1]
            if len(tks) > 3:
                if len(tks) < 12:
                    return "Tx"
                else:
                    return "Lx"

            if len(tks) == 1 and rag_tokenizer.tag(tks[0]) == "nr":
                return "Nr"

            return "Ot"

        if len(df) < 2:
            return []
        # 统计所有非首行单元格中出现最多的数据类型
        max_type = Counter([blockType(str(df.iloc[i, j])) for i in range(
            1, len(df)) for j in range(len(df.iloc[i, :]))])
        max_type = max(max_type.items(), key=lambda x: x[1])[0]

        colnm = len(df.iloc[0, :])
        hdrows = [0]  # 表头行索引列表，默认首行为表头

        # 如果表格主体是数字类型，重新识别表头
        if max_type == "Nu":
            for r in range(1, len(df)):
                tys = Counter([blockType(str(df.iloc[r, j]))
                              for j in range(len(df.iloc[r, :]))])
                tys = max(tys.items(), key=lambda x: x[1])[0]
                if tys != max_type:
                    hdrows.append(r)

        lines = []
        for i in range(1, len(df)):
            if i in hdrows:
                continue
            # 找到当前行之前最近的多层表头
            hr = [r - i for r in hdrows]
            hr = [r for r in hr if r < 0]
            t = len(hr) - 1
            while t > 0:
                if hr[t] - hr[t - 1] > 1:
                    hr = hr[t:]
                    break
                t -= 1
            # 构建每列的复合表头
            headers = []
            for j in range(len(df.iloc[i, :])):
                t = []
                for h in hr:
                    x = str(df.iloc[i + h, j]).strip()
                    if x in t:
                        continue
                    t.append(x)
                t = ",".join(t)
                if t:
                    t += ": "
                headers.append(t)
            cells = []
            for j in range(len(df.iloc[i, :])):
                if not str(df.iloc[i, j]):
                    continue
                cells.append(headers[j] + str(df.iloc[i, j]))
            lines.append(";".join(cells))

        # 窄表（≤3列）将所有行合并为一个文本块；宽表每行独立
        if colnm > 3:
            return lines
        return ["\n".join(lines)]

    def __call__(self, fnm, from_page=0, to_page=MAXIMUM_PAGE_NUMBER):
        """解析 Word 文档

        处理流程：
        1. 打开 docx 文件（支持路径或二进制内容）
        2. 遍历所有段落，按页面范围过滤
        3. 检测 lastRenderedPageBreak 标记确定分页位置
        4. 提取段落文本和样式名称
        5. 提取所有表格内容

        Args:
            fnm: 文件路径或二进制内容
            from_page: 起始页码
            to_page: 结束页码（不包含）

        Returns:
            (sections, tables) 元组：
            - sections: [(文本, 样式名), ...] 段落列表
            - tables: [结构化表格文本, ...] 表格列表
        """
        self.doc = Document(fnm) if isinstance(
            fnm, str) else Document(BytesIO(fnm))
        pn = 0  # 当前解析到的页码
        secs = []  # 解析结果：(文本, 段落样式) 列表
        for p in self.doc.paragraphs:
            if pn > to_page:
                break

            runs_within_single_paragraph = []  # 当前页范围内的 run 文本
            for run in p.runs:
                if pn > to_page:
                    break
                if from_page <= pn < to_page and p.text.strip():
                    runs_within_single_paragraph.append(run.text)

                # 检测分页标记：lastRenderedPageBreak 表示在此之后内容属于下一页
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1

            secs.append(("".join(runs_within_single_paragraph), p.style.name if hasattr(p.style, 'name') else ''))

        tbls = [self.__extract_table_content(tb) for tb in self.doc.tables]
        return secs, tbls
